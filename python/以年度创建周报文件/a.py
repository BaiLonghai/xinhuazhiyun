import calendar
import datetime
import os
import shutil
import tempfile
import re
import docx

def from_template(template_path):
    # Create a new temporary directory.
    temp_dir = tempfile.TemporaryDirectory()

    # Copy the template file to the temporary directory.
    shutil.copy(template_path, temp_dir.name)

    # Open the copied template file.
    document = docx.Document(os.path.join(temp_dir.name, template_path))

    # Return the document and the temporary directory.
    return document, temp_dir

# 设置年份
year = 2020

# 设置日期格式
date_format = "%Y年%m月%d日"
mun_date_format = "%Y%m%d"

# 遍历每一周
for week in range(1, 53):
    # 获取当前周的第一天
    first_day = datetime.datetime.strptime("%s-W%s-1" % (year, week), "%Y-W%W-%w")
    # 获取当前周的最后一天
    last_day = first_day + datetime.timedelta(days=6)

    # 转换为字符串(带年月)
    first_day_str = first_day.strftime(date_format)
    last_day_str = last_day.strftime(date_format)
    # 转换为字符串(不带年月)
    mun_first_day_str = first_day.strftime(mun_date_format)
    mun_last_day_str = last_day.strftime(mun_date_format)
    week_str = str(week)

    # 创建文件名
    filename = "项目周报_%s-%s第%s周.docx" % (
        mun_first_day_str,
        mun_last_day_str,
        week
    )

    # 使用模板文件创建新文档
    document, temp_dir = from_template("template.docx")

    # 在文档中追加文本
    #document.add_paragraph(filename)

    # 替换文档中的开始日期
    for paragraph in document.paragraphs:
        paragraph.text = re.sub(r"1975年01月01日", first_day_str, paragraph.text)
        #print(first_day_str)
    # 替换文档中的结束日期
    for paragraph in document.paragraphs:
        paragraph.text = re.sub(r"1975年02月02日", last_day_str, paragraph.text)
        #print(last_day_str)
    # 替换表格中的周数
    # 遍历每一张表格
    for table in document.tables:
        # 遍历表格的每一行
        for row in table.rows:
            # 遍历行的每一个单元格
            for cell in row.cells:
                # 遍历单元格的每一段落
                for paragraph in cell.paragraphs:
                    # 替换文本
                    paragraph.text = re.sub(r"第(\d+)周$", "第" + week_str + "周", paragraph.text)
    # 保存文档
    document.save(filename)

    # Close the temporary directory when you're done.
    temp_dir.cleanup()
