import calendar
import datetime
import os

from docx import Document

# 获取当前日期
today = datetime.datetime.now()

# 获取当前是一年的第几周
week_num = today.isocalendar()[1]

# 设置日期格式
date_format = "%Y%m%d"

# 创建一个文档对象
document = Document()

# 在文档中添加文本
document.add_paragraph("项目周报_%s-%s第%s周" % (
    today.strftime(date_format),
    today.strftime(date_format),
    week_num
))

# 遍历每一天
for day in calendar.Calendar().iterweekdays():
    # 如果是星期一至星期五
    if day in range(0, 5):
        # 获取日期
        date = today + datetime.timedelta(days=day)
        # 转换为字符串
        date_str = date.strftime(date_format)
        # 在文档中添加文本
        document.add_paragraph(date_str)

# 保存文档
document.save("项目周报_%s-%s第%s周.docx" % (
    today.strftime(date_format),
    date_str,
    week_num
))
